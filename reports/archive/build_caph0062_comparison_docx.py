"""
Build CAPH0062 eDM comparison report: ADG eDM (May 2026) vs MT eDM (March 2026).
Same visual style as the topline report.
"""
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = Path("/mnt/c/Users/User/Downloads/CareConnect_Fundamentals_4_Quote_v1.0.docx")
OUT_DIR  = Path("/mnt/c/Users/User/Downloads")
OUT_NAME = "CAPH0062_eDM_Comparison_ADG_vs_MT_v2.docx"

PURPLE     = RGBColor(0x70, 0x30, 0xA0)
DARK_NAVY  = RGBColor(0x0E, 0x28, 0x41)
BLUE       = RGBColor(0x15, 0x60, 0x82)
MID_GREY   = RGBColor(0x59, 0x59, 0x59)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREEN      = RGBColor(0x19, 0x6B, 0x24)
ORANGE     = RGBColor(0xC4, 0x6A, 0x00)


def clear_body(doc):
    body = doc.element.body
    to_remove = [child for child in body if child.tag != qn('w:sectPr')]
    for el in to_remove:
        body.remove(el)
    body.insert(0, OxmlElement('w:p'))


def _set_cell_text(cell, text):
    """Replace a header-table cell's text, keeping the first run's formatting."""
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(text)


def set_header(doc):
    """Overwrite the quote-template header text box (borderless table)."""
    tbl = doc.sections[0].header.tables[0]
    _set_cell_text(tbl.rows[0].cells[0], 'CAPH0062')
    _set_cell_text(tbl.rows[1].cells[0], 'ADG eDM (May) vs MT eDM (March)')
    _set_cell_text(tbl.rows[2].cells[0], '27 May 2026')
    _set_cell_text(tbl.rows[2].cells[1], 'V1.0')


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_table_borders(tbl, color='D0D0D0'):
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')
                el.set(qn('w:color'), color)
                borders.append(el)
            tcPr.append(borders)


def add_para(doc, text, size=10, color=None, bold=False, italic=False,
             space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or MID_GREY
    run.bold = bold
    run.italic = italic
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = PURPLE
    return p


def add_comparison_table(doc, headers, rows, col_widths=None,
                         highlight_col=None):
    """
    Table with purple header row, alternating white/light-purple body rows.
    highlight_col: 0-indexed column to render in green (positive) or orange (negative).
    """
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '70309F')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(6)
        run = p.add_run(h)
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.bold = True

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F7F0FC'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Pt(6)
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            # Colour the delta column
            if highlight_col is not None and c_idx == highlight_col:
                v = str(val).strip()
                if v.startswith('+'):
                    run.font.color.rgb = GREEN
                    run.bold = True
                elif v.startswith('−') or v.startswith('-'):
                    run.font.color.rgb = ORANGE
                    run.bold = True
                else:
                    run.font.color.rgb = MID_GREY
            else:
                run.font.color.rgb = DARK_NAVY if c_idx == 0 else MID_GREY

    add_table_borders(tbl, 'D0C0E8')

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)

    doc.add_paragraph()


def pct_change(new, old):
    if old == 0:
        return '—'
    pct = round((new - old) / old * 100)
    sign = '+' if pct >= 0 else '−'
    return f'{sign}{abs(pct)}%'


def main():
    doc = Document(TEMPLATE)
    clear_body(doc)
    set_header(doc)

    # ----------------------------------------------------------------
    # Title block
    # ----------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r = p_title.add_run('eDM Head-to-Head Comparison')
    r.font.size = Pt(18)
    r.font.color.rgb = PURPLE
    r.bold = True

    add_para(doc, 'CAPH0062  |  Rectogesic Clinical Audit  |  ADG Solus eDM (May) vs MT Solus eDM (March)',
             size=11, color=DARK_NAVY, bold=True, space_after=2)
    add_para(doc,
             'ADG send: 22 May 2026 (window 22-26)   ·   MT send: 19 March 2026 (window 19-23)'
             '   ·   Both: 5 full days post-send   ·   Prepared: 27 May 2026',
             size=9, color=MID_GREY, space_after=12)

    # ----------------------------------------------------------------
    # Section 1 — Head-to-head metrics
    # ----------------------------------------------------------------
    add_heading(doc, 'Head-to-Head: Key Metrics (5-day post-send window)')

    add_comparison_table(doc,
        headers=['Metric', 'ADG eDM  (22 May)', 'MT eDM  (19 Mar)', 'ADG vs MT'],
        rows=[
            ('Site sessions',                      '577',  '317',  pct_change(577, 317)),
            ('Unique users',                       '396',  '231',  pct_change(396, 231)),
            ('Campaign-tagged sessions (eDM)',     '482',  '228',  pct_change(482, 228)),
            ('New-to-site users',                  '243',  '109',  pct_change(243, 109)),
            ('New HCP sign-ups (prod DB)',          '94',   '40',   pct_change(94, 40)),
            ('Enrolments — Pre-learning survey F81', '33',   '25',   pct_change(33, 25)),
            ('Sign-up → enrolment rate',           '35%',  '63%',  '—'),
            ('Audit drafts — Form 161 (DB)',         '9',    '2',   '—'),
            ('Engaged sessions (GA4)',             '98%',  '100%', '—'),
            ('Pages per session (GA4)',            '2.96', '3.60', '—'),
            ('Avg engagement time / session',      '3.4 min', '5.5 min', '—'),
        ],
        col_widths=[2.55, 1.5, 1.5, 0.85],
        highlight_col=3,
    )

    # ----------------------------------------------------------------
    # Section 2 — CTA breakdown combined
    # ----------------------------------------------------------------
    add_heading(doc, 'CTA Performance — Combined')
    add_para(doc, 'Sessions driven by each link in the email. Both eDMs used the same link labels.',
             size=9, color=MID_GREY, space_after=4)

    add_comparison_table(doc,
        headers=['CTA', 'Position', 'ADG Sessions', 'MT Sessions'],
        rows=[
            ('Register Now', 'Top button',    '329', '157'),
            ('Learn More',   'Mid-email',      '79',  '36'),
            ('Register Now', 'Bottom button',  '71',  '31'),
            ('Logo / general', 'Header',        '3',   '4'),
            ('TOTAL',         '',             '482', '228'),
        ],
        col_widths=[1.8, 1.4, 1.5, 1.5],
    )

    # ----------------------------------------------------------------
    # Section 3 — Funnel
    # ----------------------------------------------------------------
    add_heading(doc, 'Conversion Funnel Comparison (prod DB)')
    add_para(doc, 'Real-HCP activity only. Excludes internal Panwar Health and TBST accounts.',
             size=9, color=MID_GREY, space_after=4)

    add_comparison_table(doc,
        headers=['Step', 'ADG (May)', 'MT (March)', 'Notes'],
        rows=[
            ('New HCP sign-ups',                     '94', '40',
             'ADG Day 1: 60  ·  MT Day 1: 25'),
            ('Started MCA learning module',          '33 (35%)', '25 (63%)',
             'Conv. from sign-ups; MT converted a higher share despite lower volume'),
            ('Completed learning module post-survey','19 (20%)', '17 (43%)',
             'Submitted after working through the module'),
            ('Audit started — Form 161 (in progress)', '9 (10%)', '2 (5%)',
             'In-progress at day 5; full audit takes weeks'),
            ('Activity evaluation',                  '0', '1',
             'Final step before cert; both windows too early for volume'),
            ('Audit submitted (complete)',            '0', '0',
             'Expected over coming weeks, not at the 5-day mark'),
        ],
        col_widths=[2.05, 0.95, 0.95, 2.45],
    )

    # ----------------------------------------------------------------
    # Section 4 — Key observations
    # ----------------------------------------------------------------
    add_heading(doc, 'Key Observations')

    bullets = [
        'ADG eDM drove markedly more traffic to the site: 577 sessions / 396 users '
        'vs MT\'s 317 / 231 (+82% / +71%), and 2.4x the sign-ups (94 vs 40). In '
        'both windows the eDM was effectively the only traffic driver - UTM-tagged '
        'sessions (482 ADG / 228 MT) make up the bulk of all site traffic.',
        'MT converted a higher share of registrants through the funnel: 63% enrolled '
        '(vs 38% for ADG) and 43% reached the post-learning survey (vs 20%). The MT '
        'cohort was smaller but more intent-driven — though it has also had two extra '
        'months to progress, so the higher rates partly reflect elapsed time, not just intent.',
        'Both eDMs show a near-identical CTA split: Register Now (top) takes ~69% of '
        'campaign sessions, Learn More ~16%, Register Now (bottom) ~14%.',
        'MT engagement time per session (5.5 min) ran higher than ADG (3.4 min), '
        'reflecting a smaller, more focused cohort spending longer per visit.',
        'Both windows are still pre-certificate: 0 completed audits / certs at the 5-day '
        'mark, as expected — the audit takes weeks to finish.',
        'Baselines confirm both spikes were eDM-driven: MT prior-week was 9 campaign '
        'sessions, ADG prior-week 77 site-wide sessions.',
    ]
    for b in bullets:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.first_line_indent = Pt(-12)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        rb = p.add_run('•  ')
        rb.font.color.rgb = PURPLE
        rb.font.size = Pt(9)
        rm = p.add_run(b)
        rm.font.size = Pt(9)
        rm.font.color.rgb = MID_GREY

    doc.add_paragraph()

    # ----------------------------------------------------------------
    # Caveat
    # ----------------------------------------------------------------
    p_cav = doc.add_paragraph()
    p_cav.paragraph_format.space_before = Pt(6)
    r1 = p_cav.add_run('Data sources: ')
    r1.font.size = Pt(8)
    r1.font.color.rgb = MID_GREY
    r1.bold = True
    r2 = p_cav.add_run(
        'GA4 property 306115293 (sessions, engagement, UTM attribution), both windows 5 full days. '
        'DB counts pulled live from prod via SSH SQL on 27 May 2026 (tbstwp_users, tbstwp_frm_items), '
        'excluding @panwarhealth.com.au and @tbstdigital.com.au accounts. '
        'Form refs: 81 Pre-learning survey, 97 Post-learning survey, 161 Retrospective analysis (audit), 209 Activity evaluation. '
        'ADG baseline: 15–21 May 2026. MT baseline: 12–18 March 2026.'
    )
    r2.font.size = Pt(8)
    r2.font.color.rgb = MID_GREY

    out_path = OUT_DIR / OUT_NAME
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    main()
