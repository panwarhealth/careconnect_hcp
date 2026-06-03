"""
Build CAPH0062 eDM topline report docx.
Clones header/footer from the quote template, replaces body with results.
"""
import shutil, copy
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

TEMPLATE = Path("/mnt/c/Users/User/Downloads/CareConnect_Fundamentals_4_Quote_v1.0.docx")
OUT_DIR  = Path("/mnt/c/Users/User/Downloads")
OUT_NAME = "CAPH0062_eDM_Topline_Results_v3.docx"

# Brand colours
PURPLE     = RGBColor(0x70, 0x30, 0xA0)
DARK_NAVY  = RGBColor(0x0E, 0x28, 0x41)
BLUE       = RGBColor(0x15, 0x60, 0x82)
MID_GREY   = RGBColor(0x59, 0x59, 0x59)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREEN      = RGBColor(0x19, 0x6B, 0x24)


def clear_body(doc):
    """Remove all paragraphs and tables from the document body, leaving one empty para."""
    body = doc.element.body
    # Collect all child elements except sectPr (section properties — has margin/page size)
    to_remove = []
    for child in body:
        if child.tag != qn('w:sectPr'):
            to_remove.append(child)
    for el in to_remove:
        body.remove(el)
    # Insert a single empty paragraph at top
    p = OxmlElement('w:p')
    body.insert(0, p)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', 'auto'))
            borders.append(el)
    tcPr.append(borders)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1'] if level == 1 else doc.styles['Normal']
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = PURPLE
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_NAVY
    return p


def add_para(doc, text, size=10, color=None, bold=False, italic=False, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or MID_GREY
    run.bold = bold
    run.italic = italic
    return p


def add_table_borders(tbl, color='D0D0D0'):
    """Apply thin borders to every cell in the table via XML."""
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')
                el.set(qn('w:color'), color)
                borders.append(el)
            tcPr.append(borders)


def add_kpi_table(doc, kpis):
    """
    kpis: list of (label, value, note) tuples — rendered as a 3-col KPI table.
    """
    cols = 3
    rows_needed = -(-len(kpis) // cols)  # ceiling div
    tbl = doc.add_table(rows=rows_needed, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set equal column widths
    total_width = Inches(6.27)
    col_width = total_width / cols
    for col in tbl.columns:
        for cell in col.cells:
            cell.width = col_width

    for idx, (label, value, note) in enumerate(kpis):
        r, c = divmod(idx, cols)
        cell = tbl.cell(r, c)
        set_cell_bg(cell, 'F2F2F2')

        p_val = cell.paragraphs[0]
        p_val.paragraph_format.space_before = Pt(6)
        p_val.paragraph_format.space_after = Pt(2)
        p_val.paragraph_format.left_indent = Pt(8)
        run_val = p_val.add_run(value)
        run_val.font.size = Pt(20)
        run_val.font.color.rgb = PURPLE
        run_val.bold = True

        p_lbl = cell.add_paragraph()
        p_lbl.paragraph_format.space_before = Pt(0)
        p_lbl.paragraph_format.space_after = Pt(2)
        p_lbl.paragraph_format.left_indent = Pt(8)
        run_lbl = p_lbl.add_run(label)
        run_lbl.font.size = Pt(9)
        run_lbl.font.color.rgb = DARK_NAVY
        run_lbl.bold = True

        if note:
            p_note = cell.add_paragraph()
            p_note.paragraph_format.space_before = Pt(0)
            p_note.paragraph_format.space_after = Pt(6)
            p_note.paragraph_format.left_indent = Pt(8)
            run_note = p_note.add_run(note)
            run_note.font.size = Pt(8)
            run_note.font.color.rgb = MID_GREY
            run_note.italic = True

    add_table_borders(tbl, 'E0D0F0')

    # Fill any empty trailing cells
    for idx in range(len(kpis), rows_needed * cols):
        r, c = divmod(idx, cols)
        cell = tbl.cell(r, c)
        set_cell_bg(cell, 'FFFFFF')

    doc.add_paragraph()  # spacer


def add_data_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '70309F')  # purple header
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(6)
        run = p.add_run(h)
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.bold = True

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F7F0FC'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Pt(6)
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_NAVY

    add_table_borders(tbl, 'D0C0E8')

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)

    doc.add_paragraph()  # spacer


def main():
    doc = Document(TEMPLATE)
    clear_body(doc)

    # ----------------------------------------------------------------
    # Title block
    # ----------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r = p_title.add_run('eDM Performance — Topline Results')
    r.font.size = Pt(18)
    r.font.color.rgb = PURPLE
    r.bold = True

    add_para(doc, 'CAPH0062  |  Rectogesic Clinical Audit  |  ADG Solus eDM', size=11, color=DARK_NAVY, bold=True, space_after=2)
    add_para(doc, f'Send date: Friday 22 May 2026   ·   Data window: 5 full days (22–26 May 2026)   ·   Prepared: {date.today().strftime("%-d %B %Y")}',
             size=9, color=MID_GREY, space_after=12)

    # ----------------------------------------------------------------
    # Section 1 — Traffic
    # ----------------------------------------------------------------
    add_heading(doc, 'Traffic & Reach', level=1)

    add_kpi_table(doc, [
        ('Site sessions (post-send, 5 days)', '577', 'vs 77 in prior 7 days (~10× daily rate)'),
        ('Unique users', '396', '243 brand new to the site'),
        ('Campaign-tagged sessions', '482', 'CAPH0062 / email medium'),
        ('New HCP sign-ups (DB)', '94', 'vs 3 in prior week — ~31× baseline'),
        ('Started MCA learning module', '33', '35% of new sign-ups enrolled within 5 days'),
        ('Engaged sessions (GA4)', '98%', '2% bounce rate · 2.96 pages/session'),
        ('Active engagement time', '3.4 min', 'avg active foreground time per session'),
    ])

    # ----------------------------------------------------------------
    # Section 2 — CTA breakdown
    # ----------------------------------------------------------------
    add_heading(doc, 'CTA Performance', level=1)
    add_para(doc, 'Which link in the email drove the most traffic:', size=9, color=MID_GREY, space_after=4)

    add_data_table(doc,
        headers=['CTA', 'Position in email', 'Sessions', 'Users', 'New users'],
        rows=[
            ('Register Now', 'Top button',    '329', '219', '153'),
            ('Learn More',   'Mid-email',      '79',  '57',  '47'),
            ('Register Now', 'Bottom button',  '71',  '53',  '40'),
            ('Logo / general link', 'Header',   '3',   '3',   '3'),
            ('TOTAL',        '',              '482', '325', '243'),
        ],
        col_widths=[1.8, 1.4, 0.85, 0.75, 0.9],
    )

    # ----------------------------------------------------------------
    # Section 3 — Funnel
    # ----------------------------------------------------------------
    add_heading(doc, 'Conversion Funnel (prod DB)', level=1)
    add_para(doc, 'Real-HCP activity only. Excludes internal Panwar Health accounts.', size=9, color=MID_GREY, space_after=4)

    add_data_table(doc,
        headers=['Step', 'Users', 'Conv. from sign-up', 'Notes'],
        rows=[
            ('New HCP sign-ups',                     '94', '—',   'Day 1: 60  ·  Day 2: 15  ·  Day 3: 9  ·  Day 4: 6  ·  Day 5: 3'),
            ('Started MCA learning module',          '35', '37%', '33 of 35 were new sign-ups from the eDM'),
            ('Completed learning module post-survey','19', '20%', '54% of those who started the module'),
            ('Audit started — Form 161 (in progress)','9', '10%', 'All 9 brand-new sign-ups from eDM; course takes weeks to complete'),
            ('Audit submitted (complete)',            '0', '—',   'Expected at 5 days — full submissions will come over coming weeks'),
            ('Certificates issued',                  '0', '—',   'Awaiting full audit + activity eval submission'),
        ],
        col_widths=[2.5, 0.7, 1.2, 2.0],
    )

    # ----------------------------------------------------------------
    # Section 4 — What to watch
    # ----------------------------------------------------------------
    add_heading(doc, 'What to Watch (Next 3–4 Weeks)', level=1)

    bullets = [
        'MCA enrolments (Form 81): 35 at day 5 — expect this to climb as new sign-ups work through the learning module.',
        'Audit submissions (Form 161): 9 in-progress drafts — full submissions will arrive over coming weeks as HCPs complete patient records.',
        'Email long-tail: eDM response decays at ~half-life 2 days; expect ~20–30 additional sessions per day through early June.',
        'Maria\'s approval queue: once HCPs submit completed audits, Maria will need to review and manually issue certs. Current backlog from this send: 0 fully submitted audits yet.',
    ]
    for b in bullets:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.first_line_indent = Pt(-12)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run_bullet = p.add_run('•  ')
        run_bullet.font.color.rgb = PURPLE
        run_bullet.font.size = Pt(9)
        run_body = p.add_run(b)
        run_body.font.size = Pt(9)
        run_body.font.color.rgb = MID_GREY

    doc.add_paragraph()

    # ----------------------------------------------------------------
    # Caveat
    # ----------------------------------------------------------------
    p_cav = doc.add_paragraph()
    p_cav.paragraph_format.space_before = Pt(8)
    r1 = p_cav.add_run('Data sources: ')
    r1.font.size = Pt(8)
    r1.font.color.rgb = MID_GREY
    r1.bold = True
    r2 = p_cav.add_run(
        'GA4 property 306115293 (sessions, engagement, UTM attribution). '
        'Prod DB via direct SQL (tbstwp_users, tbstwp_frm_items, tbstwp_usermeta, tbstwp_rcp_memberships). '
        'DB counts exclude @panwarhealth.com.au and @tbstdigital.com.au accounts. '
        'GA4 pre-week baseline: 15–21 May 2026 (no marketing sends active — organic floor).'
    )
    r2.font.size = Pt(8)
    r2.font.color.rgb = MID_GREY

    # ----------------------------------------------------------------
    # Save
    # ----------------------------------------------------------------
    out_path = OUT_DIR / OUT_NAME
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    main()
